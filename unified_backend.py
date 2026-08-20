"""
╔═══════════════════════════════════════════════════════════════════════════════╗
║         الملف المتكامل للوظائف الأساسية - مركز سرعة انجاز                  ║
║              يشمل: الإرسال والمراقبة، رسائلي، الانضمام المتقدم،             ║
║              البحث في روابطي، روابطي المحفوظة، الأدوات الأكاديمية          ║
╚═══════════════════════════════════════════════════════════════════════════════╝
"""

# ══════════════════════════════════════════════════════════════════════════════
# 1. الاستيرادات والإعدادات الأساسية
# ══════════════════════════════════════════════════════════════════════════════

import os
import json
import uuid
import time
import logging
import asyncio
import threading
import queue
import re
import random
import string
import io
import base64
import tempfile
import hashlib
import secrets
import socket
import subprocess
from datetime import datetime, timedelta
from threading import Lock, Event, Thread
from pathlib import Path
from collections import deque

# ── استيرادات Flask و Socket.IO ──
from flask import Flask, session, request, render_template, jsonify, redirect, send_file, abort, make_response, send_from_directory, Response
from flask_socketio import SocketIO, emit, join_room, leave_room

# ── استيرادات Telethon ──
from telethon import TelegramClient, events, functions
from telethon.errors import (
    SessionPasswordNeededError, PhoneCodeExpiredError, PhoneCodeInvalidError,
    PasswordHashInvalidError, FloodWaitError, UserAlreadyParticipantError,
    InviteHashExpiredError, InviteHashInvalidError
)
from telethon.sessions import StringSession
from telethon.tl.functions.channels import JoinChannelRequest, GetParticipantRequest
from telethon.tl.functions.messages import ImportChatInviteRequest, SearchGlobalRequest
from telethon.tl.functions.contacts import SearchRequest, ResolveUsernameRequest

# ── استيرادات معالجة الملفات والتحليل ──
try:
    import docx
    import pdfplumber
    import fitz  # PyMuPDF
except ImportError:
    docx = pdfplumber = fitz = None

try:
    import pandas as pd
    import numpy as np
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
    from scipy import stats
    _DATA_SCIENCE_AVAILABLE = True
except ImportError:
    _DATA_SCIENCE_AVAILABLE = False

# ── استيرادات أخرى ──
from groq import Groq
import requests
from bs4 import BeautifulSoup

# ── إعدادات التطبيق ──
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", os.urandom(24))
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# ── إعدادات السجلات ──
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ── إعدادات API ──
API_ID = os.environ.get("TELEGRAM_API_ID", "22043994")
API_HASH = os.environ.get("TELEGRAM_API_HASH", "56f64582b363d367280db96586b97801")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

# ── إعدادات GitHub ──
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN", "")
GITHUB_REPO = os.environ.get("GITHUB_REPO", "anwer1230/-Anwer_program")
GITHUB_BRANCH = os.environ.get("GITHUB_BRANCH", "main")

# ── مجلدات البيانات ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
SESSIONS_DIR = os.path.join(BASE_DIR, "sessions")
OUTPUTS_DIR = os.path.join(BASE_DIR, "pptx_app", "outputs")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(SESSIONS_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)

# ── المتغيرات العامة ──
USERS = {}
USERS_LOCK = Lock()
PREDEFINED_USERS = {}  # سيتم تحميلها ديناميكياً
_OSThread = threading.Thread

# ══════════════════════════════════════════════════════════════════════════════
# 2. دوال مساعدة مشتركة
# ══════════════════════════════════════════════════════════════════════════════

def load_settings(user_id):
    """تحميل إعدادات المستخدم من الملف المحلي"""
    user_dir = os.path.join(SESSIONS_DIR, str(user_id))
    path = os.path.join(user_dir, "settings.json")
    try:
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        logger.error(f"Error loading settings for {user_id}: {e}")
    return {}

def save_settings(user_id, settings, force=False):
    """حفظ إعدادات المستخدم في الملف المحلي"""
    try:
        if not force:
            existing = load_settings(user_id)
            if existing == settings:
                return True
        user_dir = os.path.join(SESSIONS_DIR, str(user_id))
        os.makedirs(user_dir, exist_ok=True)
        path = os.path.join(user_dir, "settings.json")
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        return True
    except Exception as e:
        logger.error(f"Error saving settings for {user_id}: {e}")
        return False

def load_string_session(user_id):
    """تحميل سلسلة الجلسة من الملف النصي"""
    path = os.path.join(SESSIONS_DIR, f"{user_id}_string.txt")
    try:
        if os.path.exists(path):
            with open(path, 'r') as f:
                return f.read().strip()
    except Exception as e:
        logger.error(f"Failed to load StringSession for {user_id}: {e}")
    return None

def save_string_session(user_id, session_str):
    """حفظ سلسلة الجلسة في ملف نصي"""
    try:
        path = os.path.join(SESSIONS_DIR, f"{user_id}_string.txt")
        with open(path, 'w') as f:
            f.write(session_str)
        return True
    except Exception as e:
        logger.error(f"Failed to save StringSession for {user_id}: {e}")
        return False

def dedupe_groups(groups):
    """إزالة التكرارات من قائمة المجموعات"""
    seen = set()
    result = []
    if isinstance(groups, str):
        groups = [g for g in groups.replace('\n', ',').split(',') if g.strip()]
    for g in groups or []:
        if not g:
            continue
        original = g.strip()
        if not original:
            continue
        norm = original.lower()
        norm = norm.replace('https://telegram.me/', 'https://t.me/')
        norm = norm.replace('http://telegram.me/', 'https://t.me/')
        norm = norm.replace('http://t.me/', 'https://t.me/')
        if '?' in norm:
            norm = norm.split('?', 1)[0]
        if '#' in norm:
            norm = norm.split('#', 1)[0]
        norm = norm.rstrip('/').strip()
        if not norm:
            continue
        if norm in seen:
            continue
        seen.add(norm)
        result.append(original)
    return result

def _clean_group_entry(raw: str) -> str:
    """تنظيف مدخل المجموعة"""
    cleaned = raw.strip()
    if not cleaned:
        return ''
    import re
    # حماية الروابط والمعرفات الرقمية
    if (cleaned.startswith('https://') or cleaned.startswith('http://')
            or cleaned.startswith('@') or re.match(r'^-?\d+$', cleaned)):
        return cleaned
    cleaned = re.sub(r'^[\s\u00b7\u2022\u25cf\u25aa\u25ab\u25fe\u25fd\u2023\u203b\u2043\u2219\*\-\–\—\.\،\,\#\>\|•●◾◾✓✦①②③④⑤⑥⑦⑧⑨⑩\d]+[\s.،:]*', '', cleaned)
    return cleaned.strip()

def extract_telegram_links(text):
    """استخراج روابط تيليجرام من النص"""
    if not text:
        return []
    patterns = [
        r'https?://t\.me/([a-zA-Z0-9_]+)(?:/\d+)?',
        r'https?://telegram\.me/([a-zA-Z0-9_]+)(?:/\d+)?',
        r'https?://t\.me/\+([a-zA-Z0-9_\-]+)',
        r'https?://telegram\.me/\+([a-zA-Z0-9_\-]+)',
        r't\.me/([a-zA-Z0-9_]+)',
        r't\.me/\+([a-zA-Z0-9_\-]+)',
        r'telegram\.me/([a-zA-Z0-9_]+)',
        r'@([a-zA-Z0-9_]{5,})',
    ]
    found_links = set()
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            clean_match = match if isinstance(match, str) else match[0] if match else ''
            if pattern.startswith(r'@'):
                clean_link = f"https://t.me/{clean_match}"
            elif '+' in clean_match or pattern.find(r'\+') != -1:
                clean_link = f"https://t.me/+{clean_match.replace('+', '')}"
            elif clean_match and not clean_match.startswith('http'):
                clean_link = f"https://t.me/{clean_match}"
            else:
                clean_link = clean_match
            if clean_link and len(clean_link) > 15:
                clean_link = clean_link.split('?')[0].split('#')[0]
                found_links.add(clean_link)
    return [{'url': link, 'username': link.split('/')[-1]} for link in sorted(found_links)]

def upload_to_github(file_path, content_bytes, commit_message="تحديث تلقائي"):
    """رفع ملف إلى GitHub"""
    if not GITHUB_TOKEN:
        return False
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{file_path}"
    headers = {"Authorization": f"token {GITHUB_TOKEN}", "Accept": "application/vnd.github.v3+json"}
    content_b64 = base64.b64encode(content_bytes).decode('utf-8')
    sha = None
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code == 200:
            sha = resp.json().get('sha')
    except Exception:
        pass
    data = {"message": commit_message, "content": content_b64, "branch": GITHUB_BRANCH}
    if sha:
        data["sha"] = sha
    try:
        r = requests.put(url, headers=headers, json=data, timeout=20)
        return r.status_code in (200, 201)
    except Exception:
        return False

def download_from_github(file_path):
    """تحميل ملف من GitHub"""
    if not GITHUB_TOKEN:
        return None
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{file_path}"
    headers = {"Authorization": f"token {GITHUB_TOKEN}", "Accept": "application/vnd.github.v3+json"}
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code == 200:
            content = resp.json().get('content', '').replace('\n', '')
            if content:
                return base64.b64decode(content)
    except Exception:
        pass
    return None

# ══════════════════════════════════════════════════════════════════════════════
# 3. الإرسال والمراقبة (Send & Monitoring)
# ══════════════════════════════════════════════════════════════════════════════

class TelegramClientManager:
    """مدير عميل تيليجرام لكل مستخدم"""
    def __init__(self, user_id):
        self.user_id = user_id
        self.client = None
        self.loop = None
        self.thread = None
        self.stop_flag = Event()
        self.is_ready = Event()
        self.monitored_keywords = []
        self.monitored_groups = []
        self._processed_msg_ids = set()

    def start_client_thread(self):
        if self.thread and self.thread.is_alive():
            return
        self.stop_flag.clear()
        self.is_ready.clear()
        self.thread = _OSThread(target=self._run_client_loop, daemon=True)
        self.thread.start()
        self.is_ready.wait(timeout=60)

    def _run_client_loop(self):
        try:
            self.loop = asyncio.new_event_loop()
            if API_ID and API_HASH:
                saved_str = load_string_session(self.user_id)
                self.client = TelegramClient(StringSession(saved_str or ''), int(API_ID), API_HASH)
            else:
                return
            self.loop.run_until_complete(self._client_main())
        except Exception as e:
            logger.error(f"Client thread error for {self.user_id}: {e}")

    async def _client_main(self):
        try:
            await self.client.connect()
            self.is_ready.set()
            await self._register_event_handlers()
            while not self.stop_flag.is_set():
                try:
                    if await self.client.is_user_authorized():
                        await self.client.run_until_disconnected()
                    else:
                        await asyncio.sleep(1)
                except Exception:
                    await asyncio.sleep(1)
        except Exception as e:
            logger.error(f"Client main error: {e}")
        finally:
            if self.client and self.client.is_connected():
                await self.client.disconnect()

    async def _register_event_handlers(self):
        if not self.client:
            return
        @self.client.on(events.NewMessage())
        async def handler(event):
            await self._handle_new_message(event)
        self.event_handlers_registered = True

    async def _handle_new_message(self, event):
        try:
            message = event.message
            if not message or not message.text:
                return
            text = message.text
            chat = await event.get_chat()
            chat_title = getattr(chat, 'title', None) or getattr(chat, 'first_name', '') or str(chat.id)
            is_outgoing = getattr(message, 'out', False)
            if not is_outgoing:
                await self._handle_auto_reply(event, message, chat_title)
            kw_list = self.monitored_keywords
            if not kw_list:
                return
            msg_uid = f"{event.chat_id}_{message.id}"
            if msg_uid in self._processed_msg_ids:
                return
            if len(self._processed_msg_ids) > 500:
                self._processed_msg_ids.clear()
            self._processed_msg_ids.add(msg_uid)
            text_clean = text.lower()
            matched = [kw for kw in kw_list if kw.lower() in text_clean]
            if matched:
                await self._trigger_keyword_alert(message, matched[0], chat_title, None, event)
        except Exception as e:
            logger.error(f"Error handling new message: {e}")

    async def _handle_auto_reply(self, event, message, chat_title):
        settings = load_settings(self.user_id)
        if not settings.get('auto_reply_enabled', True):
            return
        rules = settings.get('auto_replies', []) or []
        if not rules:
            return
        text = message.text or ''
        is_private = bool(event.is_private)
        is_group = bool(event.is_group or event.is_channel)
        for rule in rules:
            keyword = rule.get('keyword', '').strip()
            reply = rule.get('reply', '').strip()
            if not keyword or not reply:
                continue
            scope = rule.get('scope', 'all').lower()
            if scope == 'private' and not is_private:
                continue
            if scope == 'groups' and not is_group:
                continue
            match = rule.get('match', 'contains').lower()
            matched = False
            if match == 'exact':
                matched = (text.strip().lower() == keyword.lower())
            elif match == 'regex':
                matched = bool(re.search(keyword, text, re.IGNORECASE))
            else:
                matched = (keyword.lower() in text.lower())
            if matched:
                await self.client.send_message(event.chat_id, reply, reply_to=message.id)
                rule['used_count'] = int(rule.get('used_count', 0)) + 1
                rule['last_used'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                settings['auto_replies'] = rules
                save_settings(self.user_id, settings)
                break

    async def _trigger_keyword_alert(self, message, keyword, chat_title, chat_link, event):
        sender = await event.get_sender()
        sender_name = getattr(sender, 'first_name', '') or getattr(sender, 'username', '') or 'غير معروف'
        msg_text = message.text or ''
        alert_data = {
            'keyword': keyword,
            'group': chat_title,
            'sender': sender_name,
            'message': msg_text[:200],
            'full_message': msg_text,
            'timestamp': datetime.now().strftime('%H:%M:%S')
        }
        try:
            await self.client.send_message('me', f"🚨 تنبيه: {keyword}\nفي {chat_title}\nمن {sender_name}\n{msg_text[:100]}")
        except Exception:
            pass
        socketio.emit('new_alert', alert_data, to=self.user_id)

    def update_monitoring_settings(self, keywords, groups):
        self.monitored_keywords = [k.strip() for k in keywords if k.strip()]
        self.monitored_groups = [g.strip() for g in groups if g.strip()]

    def run_coroutine(self, coro):
        if not self.loop or not self.loop.is_running():
            self.start_client_thread()
        future = asyncio.run_coroutine_threadsafe(coro, self.loop)
        return future.result(timeout=60)

    async def is_group_protected(self, entity):
        """فحص حماية المجموعة"""
        return False, None

    def stop(self):
        self.stop_flag.set()
        if self.client and self.client.is_connected():
            try:
                asyncio.run_coroutine_threadsafe(self.client.disconnect(), self.loop).result(timeout=2)
            except Exception:
                pass
        if self.thread and self.thread.is_alive():
            self.thread.join(timeout=3)

class TelegramManager:
    """مدير عمليات تيليجرام الرئيسي"""
    def __init__(self):
        self.client_managers = {}
        self._smart_running = set()

    def get_client_manager(self, user_id):
        if user_id not in self.client_managers:
            self.client_managers[user_id] = TelegramClientManager(user_id)
        return self.client_managers[user_id]

    def ensure_client_active(self, user_id):
        cm = self.get_client_manager(user_id)
        cm.start_client_thread()
        return cm.is_ready.is_set()

    def send_message_async(self, user_id, entity, message, forced_action=None):
        cm = self.get_client_manager(user_id)
        if not cm.client:
            raise Exception("العميل غير متصل")
        entity_obj = cm.run_coroutine(cm.client.get_entity(entity))
        result = cm.run_coroutine(cm.client.send_message(entity_obj, message))
        return {"success": True, "message_id": result.id}

    def send_media_async(self, user_id, entity, image_files):
        cm = self.get_client_manager(user_id)
        if not cm.client:
            raise Exception("العميل غير متصل")
        entity_obj = cm.run_coroutine(cm.client.get_entity(entity))
        paths = [f['path'] for f in image_files if os.path.exists(f.get('path', ''))]
        if not paths:
            raise Exception("لا توجد صور صالحة")
        result = cm.run_coroutine(cm.client.send_file(entity_obj, paths))
        return {"success": True, "message_ids": [result.id] if not isinstance(result, list) else [r.id for r in result]}

    def send_message_with_media_async(self, user_id, entity, message, image_files):
        cm = self.get_client_manager(user_id)
        if not cm.client:
            raise Exception("العميل غير متصل")
        entity_obj = cm.run_coroutine(cm.client.get_entity(entity))
        paths = [f['path'] for f in image_files if os.path.exists(f.get('path', ''))]
        if paths:
            result = cm.run_coroutine(cm.client.send_file(entity_obj, paths, caption=message))
            return {"success": True, "message_ids": [result.id] if not isinstance(result, list) else [r.id for r in result]}
        else:
            result = cm.run_coroutine(cm.client.send_message(entity_obj, message))
            return {"success": True, "message_id": result.id}

telegram_manager = TelegramManager()

# ── دوال المراقبة والإرسال المجدول ──

def monitoring_worker(user_id):
    """خيط المراقبة والإرسال المجدول"""
    logger.info(f"Starting monitoring worker for {user_id}")
    try:
        with USERS_LOCK:
            if user_id not in USERS:
                return
            USERS[user_id]['monitoring_active'] = True
            client_manager = USERS[user_id].get('client_manager')
            settings = USERS[user_id]['settings']
        if not client_manager:
            return
        watch_words = settings.get('watch_words', [])
        send_groups = settings.get('groups', [])
        client_manager.update_monitoring_settings(watch_words, send_groups)
        _sched_dur = int(settings.get('schedule_duration', 0))
        _sched_start = time.time()
        while True:
            with USERS_LOCK:
                if user_id not in USERS or not USERS[user_id].get('is_running', False):
                    break
                user_data = USERS[user_id].copy()
            if _sched_dur > 0 and time.time() - _sched_start >= _sched_dur:
                with USERS_LOCK:
                    if user_id in USERS:
                        USERS[user_id]['is_running'] = False
                socketio.emit('schedule_status', {'running': False, 'stopped_by_duration': True}, to=user_id)
                break
            settings = user_data.get('settings', {})
            if settings.get('send_type') == 'scheduled':
                interval = int(settings.get('interval_seconds', 3600))
                last_send = user_data.get('last_scheduled_send', 0)
                if time.time() - last_send >= interval:
                    with USERS_LOCK:
                        if user_id in USERS:
                            USERS[user_id]['last_scheduled_send'] = time.time()
                    execute_scheduled_messages(user_id, settings)
            time.sleep(10)
    except Exception as e:
        logger.error(f"Monitoring worker error: {e}")
    finally:
        with USERS_LOCK:
            if user_id in USERS:
                USERS[user_id]['monitoring_active'] = False
                USERS[user_id]['is_running'] = False

def execute_scheduled_messages(user_id, settings):
    """تنفيذ الإرسال المجدول"""
    groups = settings.get('groups', [])
    message = settings.get('message', '')
    if not groups or not message:
        return
    for group in groups:
        try:
            telegram_manager.send_message_async(user_id, group, message)
            with USERS_LOCK:
                if user_id in USERS:
                    USERS[user_id]['stats']['sent'] += 1
            time.sleep(3)
        except Exception as e:
            logger.error(f"Scheduled send error to {group}: {e}")
            with USERS_LOCK:
                if user_id in USERS:
                    USERS[user_id]['stats']['errors'] += 1

# ══════════════════════════════════════════════════════════════════════════════
# 4. رسائلي (Sent Batches)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/sent_batches", methods=["GET"])
def api_sent_batches():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    with USERS_LOCK:
        batches = list(USERS[user_id].get('sent_batches', []))
    result = []
    for b in reversed(batches):
        result.append({
            "id": b["id"],
            "text": b["text"],
            "has_media": b.get("has_media", False),
            "sent_at": b["sent_at"],
            "edited_at": b.get("edited_at"),
            "sent_count": b.get("sent_count", len(b["entries"])),
            "group_count": len(b["entries"]),
        })
    return jsonify({"success": True, "batches": result})

@app.route("/api/edit_batch", methods=["POST"])
def api_edit_batch():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    data = request.json or {}
    batch_id = data.get("batch_id", "")
    new_text = data.get("new_text", "")
    if not batch_id or not new_text:
        return jsonify({"success": False, "message": "بيانات ناقصة"})
    with USERS_LOCK:
        client_manager = USERS[user_id].get('client_manager')
    if not client_manager:
        return jsonify({"success": False, "message": "يجب تسجيل الدخول أولاً"})
    def run_edit():
        try:
            client_manager.run_coroutine(_edit_batch_messages(client_manager, batch_id, new_text, user_id))
        except Exception as e:
            socketio.emit('log_update', {"message": f"❌ خطأ في التعديل: {str(e)[:100]}"}, to=user_id)
    _OSThread(target=run_edit, daemon=True).start()
    return jsonify({"success": True, "message": "⏳ جارٍ تعديل الرسائل..."})

async def _edit_batch_messages(cm, batch_id, new_text, user_id):
    with USERS_LOCK:
        ud = USERS.get(user_id, {})
        batch = next((b for b in ud.get('sent_batches', []) if b["id"] == batch_id), None)
    if not batch:
        return
    for entry in batch["entries"]:
        try:
            entity = await cm.client.get_entity(entry["group"])
            await cm.client.edit_message(entity, entry["msg_id"], new_text)
            await asyncio.sleep(0.5)
        except Exception:
            pass
    with USERS_LOCK:
        ud = USERS.get(user_id, {})
        for b in ud.get('sent_batches', []):
            if b["id"] == batch_id:
                b["text"] = new_text
                b["edited_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                break

@app.route("/api/delete_batch", methods=["POST"])
def api_delete_batch():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    data = request.json or {}
    batch_id = data.get("batch_id", "")
    if not batch_id:
        return jsonify({"success": False, "message": "batch_id مطلوب"})
    with USERS_LOCK:
        client_manager = USERS[user_id].get('client_manager')
    if not client_manager:
        return jsonify({"success": False, "message": "يجب تسجيل الدخول أولاً"})
    def run_delete():
        try:
            client_manager.run_coroutine(_delete_batch_messages(client_manager, batch_id, user_id))
        except Exception as e:
            socketio.emit('log_update', {"message": f"❌ خطأ في الحذف: {str(e)[:100]}"}, to=user_id)
    _OSThread(target=run_delete, daemon=True).start()
    return jsonify({"success": True, "message": "⏳ جارٍ حذف الرسائل..."})

async def _delete_batch_messages(cm, batch_id, user_id):
    with USERS_LOCK:
        ud = USERS.get(user_id, {})
        batch = next((b for b in ud.get('sent_batches', []) if b["id"] == batch_id), None)
    if not batch:
        return
    for entry in batch["entries"]:
        try:
            entity = await cm.client.get_entity(entry["group"])
            await cm.client.delete_messages(entity, [entry["msg_id"]])
            await asyncio.sleep(0.5)
        except Exception:
            pass
    with USERS_LOCK:
        ud = USERS.get(user_id, {})
        if ud:
            ud['sent_batches'] = [b for b in ud.get('sent_batches', []) if b["id"] != batch_id]

# ══════════════════════════════════════════════════════════════════════════════
# 5. الانضمام التلقائي المتقدم (Auto Join Advanced)
# ══════════════════════════════════════════════════════════════════════════════

async def join_telegram_group(client, group_link, user_id=None, client_manager=None):
    try:
        if group_link.startswith('https://t.me/'):
            identifier = group_link.replace('https://t.me/', '')
        elif group_link.startswith('@'):
            identifier = group_link[1:]
        else:
            identifier = group_link
        entity = await client.get_entity(identifier)
        if hasattr(entity, 'megagroup') or hasattr(entity, 'broadcast'):
            await client(JoinChannelRequest(entity))
            return {"success": True, "already_joined": False, "message": "تم الانضمام بنجاح"}
        else:
            # محاولة دعوة خاصة
            if '+' in identifier:
                result = await client(ImportChatInviteRequest(identifier))
                return {"success": True, "already_joined": False, "message": "تم الانضمام عبر رابط دعوة"}
            else:
                return {"success": False, "message": "رابط غير صالح للانضمام"}
    except UserAlreadyParticipantError:
        return {"success": True, "already_joined": True, "message": "منضم مسبقاً"}
    except FloodWaitError as e:
        return {"success": False, "message": f"انتظر {e.seconds} ثانية"}
    except Exception as e:
        return {"success": False, "message": f"فشل الانضمام: {str(e)}"}

def extract_group_names_from_text(text):
    """استخراج أسماء مجموعات محتملة من النص"""
    names = []
    patterns = [
        r'(?:قروبات|مجموعات|الدراسات العليا)\s+(جامعة\s+[^\n,]+)',
        r'([^\n,]+)\s+(?:قروب|مجموعة)',
        r'(جامعة\s+[^\n,]+)',
        r'([^\n]{5,50})',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            name = m.strip()
            if name and len(name) > 5 and len(name) < 80 and not re.match(r'^[\d\s\W]+$', name):
                names.append(name)
    return names[:20]

async def search_and_join_groups_by_name(client, names, limit=3):
    """البحث عن مجموعات بأسمائها"""
    results = []
    for name in names:
        try:
            sr = await client(SearchRequest(q=name, limit=limit))
            if sr.chats:
                for chat in sr.chats:
                    if hasattr(chat, 'username') and chat.username:
                        results.append({'name': name, 'found': f"https://t.me/{chat.username}"})
                        break
        except Exception:
            pass
    return results

def fetch_telegram_links_from_url(url, max_links=30):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.get(url, timeout=12, headers=headers)
        if resp.status_code == 200:
            soup = BeautifulSoup(resp.text, 'html.parser')
            text = soup.get_text(separator=' ', strip=True)
            extracted = extract_telegram_links(text)
            links = [link['url'] for link in extracted]
            # روابط من href
            for a in soup.find_all('a', href=True):
                href = a['href'].strip()
                if 't.me' in href and href not in links:
                    links.append(href)
            return links[:max_links]
    except Exception:
        pass
    return []

@app.route("/api/auto_join/advanced", methods=["POST"])
def api_auto_join_advanced():
    user_id = session.get('user_id', 'user_1')
    data = request.json or {}
    raw_input = data.get('links', '')
    fetch_external = data.get('fetch_external', True)
    search_by_name = data.get('search_by_name', True)
    delay = max(1, int(data.get('delay', 3)))
    max_retries = max(1, int(data.get('max_retries', 3)))

    # استخراج الروابط
    if isinstance(raw_input, str):
        raw_text = raw_input
    elif isinstance(raw_input, list):
        raw_text = "\n".join([str(item) if not isinstance(item, dict) else item.get('url', '') for item in raw_input])
    else:
        raw_text = str(raw_input)

    extracted = extract_telegram_links(raw_text)
    direct_links = [link['url'] for link in extracted]

    external_links = []
    if fetch_external:
        external_urls = re.findall(r'https?://[^\s<>"\'()]+', raw_text)
        for url in external_urls:
            if 't.me' not in url and 'telegram.me' not in url:
                fetched = fetch_telegram_links_from_url(url, 30)
                external_links.extend(fetched)

    all_raw_links = direct_links + external_links

    if not all_raw_links and search_by_name:
        group_names = extract_group_names_from_text(raw_text)
        if group_names:
            with USERS_LOCK:
                cm = USERS.get(user_id, {}).get('client_manager')
            if cm and cm.client:
                found = cm.run_coroutine(search_and_join_groups_by_name(cm.client, group_names))
                all_raw_links = [g['found'] for g in found]

    if not all_raw_links:
        return jsonify({'success': False, 'message': 'لم يتم العثور على روابط'})

    clean_links = dedupe_groups(all_raw_links)
    if not clean_links:
        return jsonify({'success': False, 'message': 'لا توجد روابط صالحة'})

    with USERS_LOCK:
        cm = USERS.get(user_id, {}).get('client_manager')
        if not cm or not cm.client:
            return jsonify({'success': False, 'message': 'العميل غير متصل'})

    # فحص العضوية المسبقة
    pending = []
    already = []
    for link in clean_links:
        try:
            entity = cm.run_coroutine(cm.client.get_entity(link))
            cm.run_coroutine(cm.client.get_participants(entity, limit=1))
            already.append(link)
        except Exception:
            pending.append(link)

    # بدء الانضمام في خيط
    def worker():
        results = {'success': 0, 'fail': 0, 'already': len(already), 'items': []}
        for idx, link in enumerate(pending, 1):
            success = False
            for attempt in range(max_retries):
                try:
                    res = cm.run_coroutine(join_telegram_group(cm.client, link, user_id, cm))
                    if res.get('success'):
                        success = True
                        results['success'] += 1
                        break
                    elif 'FloodWait' in str(res.get('message', '')):
                        wait = int(re.search(r'(\d+)', str(res.get('message', ''))).group(1)) if re.search(r'(\d+)', str(res.get('message', ''))) else 10
                        time.sleep(wait)
                    else:
                        break
                except Exception:
                    pass
                time.sleep(delay)
            if not success:
                results['fail'] += 1
            results['items'].append({'url': link, 'status': 'success' if success else 'failed'})
            socketio.emit('auto_join_progress', {'idx': idx, 'total': len(pending), 'success': results['success'], 'fail': results['fail'], 'url': link}, to=user_id)
            time.sleep(delay)
        socketio.emit('auto_join_completed', results, to=user_id)

    _OSThread(target=worker, daemon=True).start()
    return jsonify({'success': True, 'pending': len(pending), 'already': len(already)})

# ══════════════════════════════════════════════════════════════════════════════
# 6. البحث في روابطي (Search My Links)
# ══════════════════════════════════════════════════════════════════════════════

async def search_links_in_chats(client, since_date, keyword=''):
    from datetime import timezone
    since = since_date.replace(tzinfo=timezone.utc) if since_date.tzinfo is None else since_date
    found = []
    seen = set()
    try:
        async for dialog in client.iter_dialogs():
            if not dialog.entity:
                continue
            chat_title = dialog.title or 'محادثة'
            async for msg in client.iter_messages(dialog, limit=500):
                if not msg.date:
                    continue
                if msg.date.replace(tzinfo=timezone.utc) < since:
                    break
                if not msg.text:
                    continue
                for link in extract_telegram_links(msg.text):
                    url = link.get('url')
                    if url and url not in seen:
                        seen.add(url)
                        group_name = ''
                        members = 0
                        try:
                            entity = await client.get_entity(url)
                            group_name = getattr(entity, 'title', '') or getattr(entity, 'first_name', '')
                            members = getattr(entity, 'participants_count', 0)
                        except Exception:
                            pass
                        found.append({
                            'url': url,
                            'group_name': group_name,
                            'members': members,
                            'date': msg.date.strftime('%Y-%m-%d %H:%M'),
                            'chat_title': chat_title,
                            'username': link.get('username', '')
                        })
                if len(found) >= 2000:
                    break
    except Exception:
        pass
    if keyword:
        kw = keyword.lower()
        return [f for f in found if kw in (f['url'].lower() or '') or kw in (f['group_name'].lower() or '') or kw in (f['chat_title'].lower() or '')]
    return found

@app.route("/api/search_my_links/start", methods=["POST"])
def api_search_my_links_start():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"success": False, "message": "غير مسجل"}), 401
    data = request.json or {}
    keyword = data.get('keyword', '').strip()
    depth = data.get('depth', 'medium').lower()
    depth_map = {'fast': 30, 'medium': 180, 'full': 3650}
    days = depth_map.get(depth, 180)
    with USERS_LOCK:
        cm = USERS.get(user_id, {}).get('client_manager')
    if not cm or not cm.client:
        return jsonify({"success": False, "message": "العميل غير متصل"}), 400
    since = datetime.now() - timedelta(days=days)
    def worker():
        try:
            results = cm.run_coroutine(search_links_in_chats(cm.client, since, keyword))
            socketio.emit('search_links_done', {'total': len(results), 'keyword': keyword}, to=user_id)
            for item in results:
                socketio.emit('search_link_batch', {'items': [item]}, to=user_id)
        except Exception as e:
            socketio.emit('search_links_done', {'error': str(e)}, to=user_id)
    _OSThread(target=worker, daemon=True).start()
    return jsonify({"success": True, "message": f"بدأ البحث بعمق {depth}"})

@app.route("/api/search_my_links/csv", methods=["POST"])
def api_search_my_links_csv():
    data = request.json or {}
    links = data.get('links', [])
    if not links:
        return jsonify({"success": False, "message": "لا توجد روابط"})
    import csv, io
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(['الرابط', 'اسم المجموعة', 'عدد الأعضاء', 'المصدر', 'التاريخ'])
    for l in links:
        w.writerow([l.get('url',''), l.get('group_name',''), l.get('members',''), l.get('chat_title',''), l.get('date','')])
    buf.seek(0)
    return Response(buf.getvalue().encode('utf-8-sig'), mimetype='text/csv; charset=utf-8',
                    headers={'Content-Disposition': 'attachment; filename=telegram_links.csv'})

# ══════════════════════════════════════════════════════════════════════════════
# 7. روابطي المحفوظة (Saved Links)
# ══════════════════════════════════════════════════════════════════════════════

SAVED_LINKS_FILE = os.path.join(DATA_DIR, 'saved_links.json')

def load_saved_links():
    try:
        if os.path.exists(SAVED_LINKS_FILE):
            with open(SAVED_LINKS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    try:
        content = download_from_github("data/saved_links.json")
        if content:
            data = json.loads(content.decode('utf-8'))
            with open(SAVED_LINKS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return data
    except Exception:
        pass
    return {"links": []}

def save_saved_links(data):
    try:
        with open(SAVED_LINKS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    try:
        raw = json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8')
        upload_to_github("data/saved_links.json", raw, "تحديث الروابط المحفوظة")
    except Exception:
        pass

def add_saved_link(url, title=None, category='عام', notes='', source='يدوي'):
    data = load_saved_links()
    for link in data["links"]:
        if link["url"] == url:
            return False, "الرابط موجود بالفعل"
    new_link = {
        "id": str(uuid.uuid4())[:8],
        "url": url,
        "title": title or url,
        "category": category,
        "date_saved": datetime.now().isoformat(),
        "source": source,
        "notes": notes
    }
    data["links"].append(new_link)
    save_saved_links(data)
    return True, new_link

@app.route("/saved_links")
def saved_links_page():
    return render_template('saved_links.html')

@app.route("/api/saved_links", methods=["GET"])
def api_get_saved_links():
    category = request.args.get('category')
    data = load_saved_links()
    links = data["links"]
    if category:
        links = [l for l in links if l.get('category') == category]
    categories = sorted(set(l.get('category', 'عام') for l in data["links"]))
    return jsonify({"success": True, "links": links, "categories": categories, "total": len(links)})

@app.route("/api/saved_links/add", methods=["POST"])
def api_add_saved_link():
    data = request.json or {}
    url = data.get('url', '').strip()
    title = data.get('title', '').strip() or url
    category = data.get('category', 'عام')
    notes = data.get('notes', '')
    source = data.get('source', 'يدوي')
    if not url:
        return jsonify({"success": False, "message": "الرابط مطلوب"})
    success, result = add_saved_link(url, title, category, notes, source)
    if success:
        return jsonify({"success": True, "link": result})
    return jsonify({"success": False, "message": result})

@app.route("/api/saved_links/delete", methods=["POST"])
def api_delete_saved_link():
    data = request.json or {}
    link_id = data.get('id')
    if not link_id:
        return jsonify({"success": False, "message": "معرف الرابط مطلوب"})
    data_links = load_saved_links()
    data_links["links"] = [l for l in data_links["links"] if l["id"] != link_id]
    save_saved_links(data_links)
    return jsonify({"success": True})

@app.route("/api/saved_links/add_batch", methods=["POST"])
def api_add_saved_links_batch():
    data = request.json or {}
    urls = data.get('urls', [])
    category = data.get('category', 'عام')
    source = data.get('source', 'دفعة')
    added = []
    skipped = []
    for url in urls:
        url = url.strip()
        if not url:
            continue
        success, _ = add_saved_link(url, url, category, '', source)
        if success:
            added.append(url)
        else:
            skipped.append(url)
    return jsonify({"success": True, "added": added, "skipped": skipped, "added_count": len(added), "skipped_count": len(skipped)})

@app.route("/api/saved_links/send_to_auto_join", methods=["POST"])
def api_send_saved_to_auto_join():
    data = request.json or {}
    link_ids = data.get('ids', [])
    data_links = load_saved_links()
    if link_ids:
        links = [l for l in data_links["links"] if l["id"] in link_ids]
    else:
        links = data_links["links"]
    urls = [l["url"] for l in links]
    session['auto_join_links'] = urls
    return jsonify({"success": True, "urls": urls, "count": len(urls)})

# ══════════════════════════════════════════════════════════════════════════════
# 8. الأدوات الأكاديمية (Academic Tools)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/academic")
def academic_analysis():
    return render_template('academic.html')

@app.route("/tools/analyze_stats", methods=["POST"])
def api_academic_analyze_stats():
    try:
        data = request.get_json(force=True, silent=True) or {}
        raw = data.get('data', '')
        nums = [float(x) for x in re.findall(r'[-+]?\d*\.?\d+', str(raw)) if x]
        if len(nums) < 2:
            return jsonify({"error": "أدخل على الأقل رقمين للتحليل"}), 400
        arr = np.array(nums)
        mode_val = float(stats.mode(arr, keepdims=True).mode[0])
        stats_result = {
            "count": int(len(arr)),
            "sum": round(float(np.sum(arr)), 4),
            "mean": round(float(np.mean(arr)), 4),
            "median": round(float(np.median(arr)), 4),
            "mode": round(mode_val, 4),
            "std": round(float(np.std(arr)), 4),
            "variance": round(float(np.var(arr)), 4),
            "min": round(float(np.min(arr)), 4),
            "max": round(float(np.max(arr)), 4),
            "range": round(float(np.max(arr) - np.min(arr)), 4),
            "q1": round(float(np.percentile(arr, 25)), 4),
            "q3": round(float(np.percentile(arr, 75)), 4),
            "iqr": round(float(np.percentile(arr, 75) - np.percentile(arr, 25)), 4),
            "skewness": round(float(stats.skew(arr)), 4),
            "kurtosis": round(float(stats.kurtosis(arr)), 4),
        }
        fig, axes = plt.subplots(1, 2, figsize=(10, 4))
        axes[0].hist(arr, bins='auto', color='#4e73df', edgecolor='white', alpha=0.85)
        axes[0].axvline(stats_result['mean'], color='red', linestyle='--', label=f"المتوسط: {stats_result['mean']}")
        axes[0].axvline(stats_result['median'], color='green', linestyle='--', label=f"الوسيط: {stats_result['median']}")
        axes[0].set_title('توزيع البيانات')
        axes[0].legend()
        axes[1].boxplot(arr, vert=True, patch_artist=True, boxprops=dict(facecolor='#4e73df', alpha=0.7))
        axes[1].set_title('المربع الجذري')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        chart_b64 = base64.b64encode(buf.read()).decode('utf-8')
        return jsonify({"success": True, "stats": stats_result, "chart": chart_b64})
    except Exception as e:
        logger.error(f"Analyze stats error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/tools/format_file", methods=["POST"])
def api_academic_format_file():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم رفع ملف"}), 400
        f = request.files['file']
        filename = f.filename.lower()
        content_parts = []
        file_bytes = f.read()
        if filename.endswith('.pdf'):
            if pdfplumber is None:
                return jsonify({"error": "مكتبة PDF غير متاحة"}), 500
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for i, page in enumerate(pdf.pages[:50], 1):
                    text = page.extract_text() or ''
                    if text.strip():
                        content_parts.append(f"--- صفحة {i} ---\n{text.strip()}")
                    for tbl in page.extract_tables():
                        if tbl:
                            rows_txt = '\n'.join('\t'.join(str(c or '') for c in row) for row in tbl)
                            content_parts.append(f"[جدول]\n{rows_txt}")
        elif filename.endswith('.docx'):
            if docx is None:
                return jsonify({"error": "مكتبة Word غير متاحة"}), 500
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            for para in doc_obj.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            for tbl in doc_obj.tables:
                rows_txt = '\n'.join('\t'.join(c.text for c in row.cells) for row in tbl.rows)
                content_parts.append(f"[جدول]\n{rows_txt}")
        elif filename.endswith('.txt'):
            content_parts.append(file_bytes.decode('utf-8', errors='replace'))
        else:
            return jsonify({"error": "صيغة غير مدعومة"}), 400
        full_text = '\n\n'.join(content_parts)
        ai_summary = None
        use_ai = request.form.get('use_ai', 'false').lower() == 'true'
        if use_ai and full_text.strip() and GROQ_API_KEY:
            try:
                client = Groq(api_key=GROQ_API_KEY)
                resp = client.chat.completions.create(
                    model='llama-3.3-70b-versatile',
                    messages=[{"role": "user", "content": f"لخّص هذا النص بشكل أكاديمي:\n\n{full_text[:6000]}"}],
                    max_tokens=800,
                    temperature=0.3
                )
                ai_summary = resp.choices[0].message.content
            except Exception as e:
                ai_summary = f"[تعذّر التلخيص: {e}]"
        return jsonify({"success": True, "text": full_text, "ai_summary": ai_summary, "words": len(full_text.split())})
    except Exception as e:
        logger.error(f"Format file error: {e}")
        return jsonify({"error": str(e)}), 500

# ══════════════════════════════════════════════════════════════════════════════
# 9. واجهات إضافية (Routes) – الإرسال والمراقبة
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/send_now", methods=["POST"])
def api_send_now():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    data = request.json or {}
    message = data.get('message', '').strip()
    groups = data.get('groups', '').strip()
    images = data.get('images', [])
    send_to_all = data.get('send_to_all', False)
    if not message and not images:
        return jsonify({"success": False, "message": "الرسالة أو الصورة مطلوبة"})
    with USERS_LOCK:
        cm = USERS.get(user_id, {}).get('client_manager')
        if not cm or not cm.client:
            return jsonify({"success": False, "message": "العميل غير متصل"})

    if send_to_all:
        dialogs = cm.run_coroutine(cm.client.get_dialogs())
        groups_list = []
        for d in dialogs:
            entity = d.entity
            if hasattr(entity, 'megagroup') or hasattr(entity, 'broadcast'):
                uname = getattr(entity, 'username', None)
                if uname:
                    groups_list.append(f"https://t.me/{uname}")
        if not groups_list:
            return jsonify({"success": False, "message": "لا توجد مجموعات"})
    else:
        if not groups:
            return jsonify({"success": False, "message": "المجموعات مطلوبة"})
        groups_list = dedupe_groups([g.strip() for g in groups.replace('\n', ',').split(',') if g.strip()])

    image_files = []
    if images:
        for img in images:
            try:
                raw_data = img.get('data', '')
                if ',' in raw_data:
                    b64 = raw_data.split(',', 1)[1]
                else:
                    b64 = raw_data
                img_bytes = base64.b64decode(b64)
                ext = img.get('type', 'image/jpeg').split('/')[-1].lower()
                if ext in ('jpeg', 'jpg'):
                    ext = 'jpg'
                elif ext not in ('png', 'gif', 'webp', 'bmp'):
                    ext = 'jpg'
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}', mode='wb')
                tmp.write(img_bytes)
                tmp.close()
                image_files.append({'path': tmp.name, 'name': img.get('name', f'image.{ext}'), 'type': img.get('type', 'image/jpeg')})
            except Exception as e:
                logger.error(f"Image processing error: {e}")

    def send_worker():
        successful = 0
        failed = 0
        batch_id = str(uuid.uuid4())
        batch_entries = []
        for group in groups_list:
            try:
                if images and message:
                    res = telegram_manager.send_message_with_media_async(user_id, group, message, image_files)
                elif images:
                    res = telegram_manager.send_media_async(user_id, group, image_files)
                else:
                    res = telegram_manager.send_message_async(user_id, group, message)
                if isinstance(res, dict) and res.get('success'):
                    successful += 1
                    mid = res.get('message_id') or (res.get('message_ids') or [None])[0]
                    if mid:
                        batch_entries.append({"group": group, "msg_id": mid})
                    with USERS_LOCK:
                        if user_id in USERS:
                            USERS[user_id]['stats']['sent'] += 1
                else:
                    failed += 1
                time.sleep(3)
            except Exception as e:
                failed += 1
                with USERS_LOCK:
                    if user_id in USERS:
                        USERS[user_id]['stats']['errors'] += 1
        if batch_entries:
            batch_record = {
                "id": batch_id,
                "text": message or "",
                "has_media": bool(images),
                "sent_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "sent_count": successful,
                "entries": batch_entries
            }
            with USERS_LOCK:
                ud = USERS.get(user_id)
                if ud:
                    ud.setdefault('sent_batches', []).append(batch_record)
                    if len(ud['sent_batches']) > 100:
                        ud['sent_batches'] = ud['sent_batches'][-100:]
                    try:
                        settings = load_settings(user_id)
                        settings['sent_batches'] = ud['sent_batches']
                        save_settings(user_id, settings)
                    except Exception:
                        pass
            socketio.emit('batch_saved', batch_record, to=user_id)
        for img_file in image_files:
            try:
                if os.path.exists(img_file['path']):
                    os.unlink(img_file['path'])
            except Exception:
                pass
        socketio.emit('log_update', {"message": f"✅ تم الإرسال: {successful} نجح, {failed} فشل"}, to=user_id)

    _OSThread(target=send_worker, daemon=True).start()
    return jsonify({"success": True, "message": f"بدء الإرسال إلى {len(groups_list)} مجموعة"})

@app.route("/api/start_monitoring", methods=["POST"])
def api_start_monitoring():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    with USERS_LOCK:
        if USERS[user_id].get('is_running', False):
            return jsonify({"success": False, "message": "المراقبة تعمل بالفعل"})
        USERS[user_id]['is_running'] = True
    _OSThread(target=monitoring_worker, args=(user_id,), daemon=True).start()
    return jsonify({"success": True, "message": "بدأت المراقبة"})

@app.route("/api/stop_monitoring", methods=["POST"])
def api_stop_monitoring():
    user_id = session.get('user_id')
    if not user_id or user_id not in USERS:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    with USERS_LOCK:
        USERS[user_id]['is_running'] = False
    return jsonify({"success": True, "message": "تم إيقاف المراقبة"})

@app.route("/api/save_settings", methods=["POST"])
def api_save_settings():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"success": False, "message": "غير مسجّل"}), 401
    data = request.json or {}
    current = load_settings(user_id)
    current.update({
        'message': data.get('message', ''),
        'groups': dedupe_groups(data.get('groups', '')),
        'interval_seconds': int(data.get('interval_seconds', 3600)),
        'watch_words': [w.strip() for w in data.get('watch_words', '').split('\n') if w.strip()],
        'send_type': data.get('send_type', 'manual'),
        'schedule_duration': int(float(data.get('schedule_duration_hours', 0) or 0) * 3600),
        'auto_reply_enabled': data.get('auto_reply_enabled', True),
    })
    if save_settings(user_id, current):
        with USERS_LOCK:
            if user_id in USERS:
                USERS[user_id]['settings'] = current
        return jsonify({"success": True, "message": "تم حفظ الإعدادات"})
    return jsonify({"success": False, "message": "فشل الحفظ"})

# ══════════════════════════════════════════════════════════════════════════════
# 10. تشغيل الخادم (للتطوير المحلي)
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    try:
        with open(os.path.join(DATA_DIR, 'dynamic_users.json'), 'r') as f:
            PREDEFINED_USERS = json.load(f)
    except:
        PREDEFINED_USERS = {}
    for uid in PREDEFINED_USERS:
        settings = load_settings(uid)
        USERS[uid] = {
            'client_manager': None,
            'settings': settings,
            'is_running': False,
            'stats': {'sent': 0, 'errors': 0},
            'connected': False,
            'authenticated': False,
            'sent_batches': settings.get('sent_batches', []),
            'monitoring_active': False,
        }
    socketio.run(app, host='0.0.0.0', port=5000, debug=False)
